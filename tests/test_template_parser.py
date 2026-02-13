from __future__ import annotations

from pathlib import Path

from docx import Document

from mrm_deepagent.models import SectionType
from mrm_deepagent.template_parser import parse_template, validate_template


def test_parse_template_extracts_sections(template_path: Path) -> None:
    parsed = parse_template(template_path)
    assert len(parsed.sections) == 4
    fill_sections = [
        section for section in parsed.sections if section.section_type == SectionType.FILL
    ]
    assert len(fill_sections) == 2
    assert "model_validated" in fill_sections[0].checkbox_tokens


def test_validate_template_detects_duplicate_ids(duplicate_template_path: Path) -> None:
    parsed = parse_template(duplicate_template_path)
    errors = validate_template(parsed)
    assert any("Duplicate section ID: data_description" in error for error in errors)


def test_parse_template_detects_malformed_marker(tmp_path: Path) -> None:
    path = tmp_path / "bad_marker.docx"
    document = Document()
    document.add_heading("[FILL][ID bad] malformed marker", level=1)
    document.add_paragraph("Body")
    document.save(path)

    parsed = parse_template(path)
    errors = validate_template(parsed)
    assert any("Malformed marker heading" in error for error in errors)
