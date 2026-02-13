from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document


def build_template_docx(
    path: Path,
    duplicate_fill_id: bool = False,
    include_table: bool = False,
    include_untagged_heading: bool = False,
) -> Path:
    document = Document()
    document.add_heading("[FILL][ID:exec_summary] Executive Summary", level=1)
    document.add_paragraph(
        "Summarize model objective and business impact. [[SECTION_CONTENT]] "
        "[[CHECK:model_validated]]"
    )
    if include_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Model validated"
        table.cell(0, 1).text = "[[CHECK:model_validated]]"
        table.cell(1, 0).text = "Validation owner"
        table.cell(1, 1).text = "Model Risk Team"

    document.add_heading("[FILL][ID:data_description] Data Description", level=1)
    document.add_paragraph("Describe training data sources and quality checks. [[SECTION_CONTENT]]")

    if duplicate_fill_id:
        document.add_heading("[FILL][ID:data_description] Duplicate ID", level=1)
        document.add_paragraph("Duplicate section to force validation failure.")

    if include_untagged_heading:
        document.add_heading("4. Model Change Management", level=1)
        document.add_paragraph("Untagged heading should be treated as a fillable section.")

    document.add_heading("[SKIP][ID:validator_notes] Validator Notes", level=1)
    document.add_paragraph("This section is for validator-only completion.")

    document.add_heading("[VALIDATOR][ID:validation_results] Validation Results", level=1)
    document.add_paragraph("Completed by model validation team.")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def build_template_markdown(path: Path, include_missing_token: bool = False) -> Path:
    model_overview_block = """
# [FILL][ID:model_overview] Model Overview

Requirements:
- Describe model objective.

Response:
[[SECTION_CONTENT]]
""".strip()
    if include_missing_token:
        model_overview_block = model_overview_block.replace("[[SECTION_CONTENT]]", "Response here.")

    text = (
        model_overview_block
        + "\n\n---\n\n"
        + """
# [FILL][ID:model_purpose] Purpose

Requirements:
- Explain business purpose.
- Confirm intended use [[CHECK:intended_use_defined]].

Response:
[[SECTION_CONTENT]]

---

# [SKIP][ID:reviewer_notes] Reviewer Notes

Reserved for reviewer.

---

# [VALIDATOR][ID:validation_signoff] Validation Signoff

Reserved for validator.
""".strip()
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text + "\n", encoding="utf-8")
    return path


@pytest.fixture
def template_path(tmp_path: Path) -> Path:
    return build_template_docx(tmp_path / "template.docx")


@pytest.fixture
def duplicate_template_path(tmp_path: Path) -> Path:
    return build_template_docx(tmp_path / "duplicate_template.docx", duplicate_fill_id=True)


@pytest.fixture
def table_template_path(tmp_path: Path) -> Path:
    return build_template_docx(tmp_path / "table_template.docx", include_table=True)


@pytest.fixture
def untagged_template_path(tmp_path: Path) -> Path:
    return build_template_docx(tmp_path / "untagged_template.docx", include_untagged_heading=True)


@pytest.fixture
def markdown_template_path(tmp_path: Path) -> Path:
    return build_template_markdown(tmp_path / "template.md")


@pytest.fixture
def markdown_template_missing_token_path(tmp_path: Path) -> Path:
    return build_template_markdown(tmp_path / "bad_template.md", include_missing_token=True)
