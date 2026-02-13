from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from mrm_deepagent.docx_applier import apply_draft_to_template
from mrm_deepagent.exceptions import AlreadyAppliedError
from mrm_deepagent.models import (
    CheckboxToken,
    DraftDocument,
    DraftSection,
    DraftStatus,
    MissingItem,
)


def _build_draft() -> DraftDocument:
    return DraftDocument(
        sections=[
            DraftSection(
                id="exec_summary",
                title="Executive Summary",
                status=DraftStatus.COMPLETE,
                checkboxes=[CheckboxToken(name="model_validated", checked=True)],
                attachments=[],
                evidence=["train.py:1"],
                missing_items=[],
                body="Filled summary text.",
            ),
            DraftSection(
                id="data_description",
                title="Data Description",
                status=DraftStatus.PARTIAL,
                checkboxes=[],
                attachments=[],
                evidence=[],
                missing_items=[
                    MissingItem(
                        id="missing_data_owner",
                        section_id="data_description",
                        question="Need owner",
                    )
                ],
                body="Partial body text.",
            ),
        ]
    )


def test_apply_updates_fill_sections(tmp_path: Path, template_path: Path) -> None:
    out_path = tmp_path / "applied.docx"
    report = apply_draft_to_template(template_path, _build_draft(), out_path)
    assert out_path.exists()
    assert report.unresolved_section_ids == ["data_description"]

    document = Document(str(out_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Filled summary text." in text
    assert "\u2612" in text
    assert "validator-only completion" in text
    assert "UNRESOLVED" in text


def test_apply_handles_table_templates(tmp_path: Path, table_template_path: Path) -> None:
    out_path = tmp_path / "out.docx"
    apply_draft_to_template(table_template_path, _build_draft(), out_path)
    document = Document(str(out_path))
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "\u2612" in table_text


def test_apply_rejects_already_applied_without_force(tmp_path: Path, template_path: Path) -> None:
    first_output = tmp_path / "first.docx"
    apply_draft_to_template(template_path, _build_draft(), first_output)
    with pytest.raises(AlreadyAppliedError):
        apply_draft_to_template(first_output, _build_draft(), tmp_path / "second.docx")


def test_apply_allows_force_on_already_applied(tmp_path: Path, template_path: Path) -> None:
    first_output = tmp_path / "first.docx"
    apply_draft_to_template(template_path, _build_draft(), first_output)
    second_output = tmp_path / "second.docx"
    apply_draft_to_template(first_output, _build_draft(), second_output, force=True)
    assert second_output.exists()
