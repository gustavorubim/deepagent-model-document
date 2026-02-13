"""Generate a realistic fictitious MRM template with mixed tagging and tables."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def build_template(output_path: Path) -> None:
    document = Document()

    document.add_heading("Model Risk Management Documentation Template", level=0)
    document.add_paragraph(
        "Fictitious internal template inspired by supervisory model risk guidance."
    )

    document.add_heading("Document Control", level=1)
    control_table = document.add_table(rows=4, cols=2)
    control_table.cell(0, 0).text = "Document owner"
    control_table.cell(0, 1).text = "[[SECTION_CONTENT]]"
    control_table.cell(1, 0).text = "Version"
    control_table.cell(1, 1).text = "0.1"
    control_table.cell(2, 0).text = "Review date"
    control_table.cell(2, 1).text = "[[SECTION_CONTENT]]"
    control_table.cell(3, 0).text = "Approved by"
    control_table.cell(3, 1).text = "[[SECTION_CONTENT]]"

    document.add_heading("[FILL][ID:purpose_scope] 1. Purpose and Scope", level=1)
    document.add_paragraph(
        "Describe intended business use, regulatory scope, and model boundaries. "
        "Include dependencies, assumptions, and known exclusions. [[SECTION_CONTENT]]"
    )

    document.add_heading("2. Model Inventory and Risk Tiering [FILL]", level=1)
    document.add_paragraph(
        "Summarize inventory metadata and inherent model risk rating rationale. [[SECTION_CONTENT]]"
    )
    tier_table = document.add_table(rows=5, cols=2)
    tier_table.cell(0, 0).text = "Model ID"
    tier_table.cell(0, 1).text = "[[SECTION_CONTENT]]"
    tier_table.cell(1, 0).text = "Model owner"
    tier_table.cell(1, 1).text = "[[SECTION_CONTENT]]"
    tier_table.cell(2, 0).text = "Risk tier: high"
    tier_table.cell(2, 1).text = "[[CHECK:risk_tier_high]]"
    tier_table.cell(3, 0).text = "Risk tier: medium"
    tier_table.cell(3, 1).text = "[[CHECK:risk_tier_medium]]"
    tier_table.cell(4, 0).text = "Risk tier: low"
    tier_table.cell(4, 1).text = "[[CHECK:risk_tier_low]]"

    document.add_heading("[ID:model_development][FILL] 3. Development and Implementation", level=1)
    document.add_paragraph(
        "Explain data lineage, feature engineering, training pipeline, and deployment controls. "
        "[[SECTION_CONTENT]]"
    )
    dev_table = document.add_table(rows=5, cols=3)
    dev_table.cell(0, 0).text = "Control"
    dev_table.cell(0, 1).text = "Implemented"
    dev_table.cell(0, 2).text = "Evidence"
    dev_table.cell(1, 0).text = "Data quality checks"
    dev_table.cell(1, 1).text = "[[CHECK:data_quality_checks]]"
    dev_table.cell(1, 2).text = "[[SECTION_CONTENT]]"
    dev_table.cell(2, 0).text = "Code review"
    dev_table.cell(2, 1).text = "[[CHECK:code_review_completed]]"
    dev_table.cell(2, 2).text = "[[SECTION_CONTENT]]"
    dev_table.cell(3, 0).text = "Reproducible training"
    dev_table.cell(3, 1).text = "[[CHECK:reproducible_training]]"
    dev_table.cell(3, 2).text = "[[SECTION_CONTENT]]"
    dev_table.cell(4, 0).text = "Change log maintained"
    dev_table.cell(4, 1).text = "[[CHECK:change_log_maintained]]"
    dev_table.cell(4, 2).text = "[[SECTION_CONTENT]]"

    document.add_heading(
        "[FILL][ID:model_performance] 4. Model Performance and Limitations", level=1
    )
    document.add_paragraph(
        "Provide performance metrics, confidence intervals, limitations, and material model "
        "assumptions. [[SECTION_CONTENT]]"
    )
    perf_table = document.add_table(rows=4, cols=2)
    perf_table.cell(0, 0).text = "Metric"
    perf_table.cell(0, 1).text = "Value / Reference"
    perf_table.cell(1, 0).text = "MAE"
    perf_table.cell(1, 1).text = "[[SECTION_CONTENT]]"
    perf_table.cell(2, 0).text = "RMSE"
    perf_table.cell(2, 1).text = "[[SECTION_CONTENT]]"
    perf_table.cell(3, 0).text = "R2"
    perf_table.cell(3, 1).text = "[[SECTION_CONTENT]]"

    document.add_heading(
        "[VALIDATOR][ID:validation_independent] 5. Independent Validation", level=1
    )
    document.add_paragraph("Reserved for independent validation function.")
    val_table = document.add_table(rows=3, cols=2)
    val_table.cell(0, 0).text = "Validation status"
    val_table.cell(0, 1).text = "[[CHECK:validator_approved]]"
    val_table.cell(1, 0).text = "Validation findings summary"
    val_table.cell(1, 1).text = "To be completed by validator."
    val_table.cell(2, 0).text = "Residual model risk"
    val_table.cell(2, 1).text = "To be completed by validator."

    document.add_heading("6. Ongoing Monitoring and Change Management [FILL]", level=1)
    document.add_paragraph(
        "Describe monitoring cadence, thresholds, override process, and incident escalation. "
        "[[SECTION_CONTENT]]"
    )
    mon_table = document.add_table(rows=5, cols=3)
    mon_table.cell(0, 0).text = "Monitoring control"
    mon_table.cell(0, 1).text = "Enabled"
    mon_table.cell(0, 2).text = "Threshold / trigger"
    mon_table.cell(1, 0).text = "Performance drift checks"
    mon_table.cell(1, 1).text = "[[CHECK:monitor_drift]]"
    mon_table.cell(1, 2).text = "[[SECTION_CONTENT]]"
    mon_table.cell(2, 0).text = "Data drift checks"
    mon_table.cell(2, 1).text = "[[CHECK:monitor_data_drift]]"
    mon_table.cell(2, 2).text = "[[SECTION_CONTENT]]"
    mon_table.cell(3, 0).text = "Override tracking"
    mon_table.cell(3, 1).text = "[[CHECK:monitor_overrides]]"
    mon_table.cell(3, 2).text = "[[SECTION_CONTENT]]"
    mon_table.cell(4, 0).text = "Periodic revalidation"
    mon_table.cell(4, 1).text = "[[CHECK:periodic_revalidation]]"
    mon_table.cell(4, 2).text = "[[SECTION_CONTENT]]"

    document.add_heading("[FILL] 7. Governance and Approvals", level=1)
    document.add_paragraph(
        "Document roles, committee approvals, and exception governance. [[SECTION_CONTENT]]"
    )
    gov_table = document.add_table(rows=5, cols=3)
    gov_table.cell(0, 0).text = "Role"
    gov_table.cell(0, 1).text = "Assigned"
    gov_table.cell(0, 2).text = "Name / date"
    gov_table.cell(1, 0).text = "Model owner"
    gov_table.cell(1, 1).text = "[[CHECK:role_model_owner_assigned]]"
    gov_table.cell(1, 2).text = "[[SECTION_CONTENT]]"
    gov_table.cell(2, 0).text = "Model developer"
    gov_table.cell(2, 1).text = "[[CHECK:role_model_developer_assigned]]"
    gov_table.cell(2, 2).text = "[[SECTION_CONTENT]]"
    gov_table.cell(3, 0).text = "Model reviewer"
    gov_table.cell(3, 1).text = "[[CHECK:role_model_reviewer_assigned]]"
    gov_table.cell(3, 2).text = "[[SECTION_CONTENT]]"
    gov_table.cell(4, 0).text = "Model approver"
    gov_table.cell(4, 1).text = "[[CHECK:role_model_approver_assigned]]"
    gov_table.cell(4, 2).text = "[[SECTION_CONTENT]]"

    document.add_heading(
        "[FILL][ID:exceptions_controls] 8. Exceptions and Compensating Controls", level=1
    )
    document.add_paragraph(
        "If policy exceptions exist, provide rationale, approval, expiry date, and compensating "
        "controls. [[SECTION_CONTENT]]"
    )

    document.add_heading("[SKIP][ID:validator_internal_notes] Internal Validator Notes", level=1)
    document.add_paragraph("Reserved for validation function internal notes.")

    document.add_heading("Appendix A - Validator Completion [VALIDATOR]", level=1)
    appendix = document.add_table(rows=4, cols=2)
    appendix.cell(0, 0).text = "Final model risk rating"
    appendix.cell(0, 1).text = "Validator completion required"
    appendix.cell(1, 0).text = "Open issues count"
    appendix.cell(1, 1).text = "Validator completion required"
    appendix.cell(2, 0).text = "Approval date"
    appendix.cell(2, 1).text = "Validator completion required"
    appendix.cell(3, 0).text = "Follow-up due date"
    appendix.cell(3, 1).text = "Validator completion required"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    build_template(Path("examples/fictitious_mrm_template.docx"))


if __name__ == "__main__":
    main()
